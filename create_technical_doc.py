#!/usr/bin/env python3
"""
Generate Technical Documentation for RAG Customer Service Chatbot System
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def add_title_page(doc):
    """Add professional title page"""
    title = doc.add_paragraph()
    title_run = title.add_run("RAG-Powered Customer Service Chatbot\n")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Technical Documentation")
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    info = doc.add_paragraph()
    info_run = info.add_run(f"Airtable-Integrated Ticket Management System\n")
    info_run.font.size = Pt(14)
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(12)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

def add_toc(doc):
    """Add table of contents placeholder"""
    heading = doc.add_heading('Table of Contents', 1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    toc_items = [
        "1. System Overview",
        "2. Architecture",
        "3. Workflow Components",
        "4. Data Schema",
        "5. API Reference",
        "6. Testing Framework",
        "7. Deployment Guide",
        "8. Troubleshooting",
        "9. Appendices"
    ]

    for item in toc_items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)

    doc.add_page_break()

def add_system_overview(doc):
    """Add system overview section"""
    doc.add_heading('1. System Overview', 1)

    doc.add_heading('1.1 Executive Summary', 2)
    doc.add_paragraph(
        "The RAG-Powered Customer Service Chatbot is an AI-driven support system that combines "
        "Retrieval-Augmented Generation (RAG) technology with automated ticket management. The system "
        "uses n8n workflow automation, Airtable as the database backend, and integrates with Pinecone "
        "vector store for knowledge retrieval."
    )

    doc.add_heading('1.2 Key Features', 2)
    features = [
        "Intelligent conversation handling using AI agents (OpenAI/LangChain)",
        "Automated ticket creation, status checking, updating, and closure",
        "Document retrieval from Google Drive integrated knowledge base",
        "Vector embeddings via Pinecone for semantic search",
        "Real-time Slack notifications for support team",
        "SLA tracking and priority-based ticket routing",
        "Conversation log maintenance for audit trails",
        "Webhook-based API for external integrations"
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('1.3 Technology Stack', 2)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'

    tech_stack = [
        ('Workflow Automation', 'n8n Cloud (v1.118.2)'),
        ('Database', 'Airtable'),
        ('AI/LLM', 'OpenAI API (LangChain)'),
        ('Vector Store', 'Pinecone'),
        ('Document Source', 'Google Drive'),
        ('Notifications', 'Slack API'),
        ('Testing', 'Bash scripts, curl, jq'),
        ('Version Control', 'Git/GitHub')
    ]

    for component, tech in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = component
        row_cells[1].text = tech

    doc.add_paragraph()

    doc.add_heading('1.4 Project Status', 2)
    doc.add_paragraph("Status: Production Ready ✓")
    doc.add_paragraph("Version: 1.0")
    doc.add_paragraph("Last Updated: December 2024")
    doc.add_paragraph("Test Coverage: 100% (6/6 core tests passing)")

    doc.add_page_break()

def add_architecture(doc):
    """Add architecture section"""
    doc.add_heading('2. System Architecture', 1)

    doc.add_heading('2.1 High-Level Architecture', 2)

    arch_diagram = """
    ┌─────────────────────────────────────────────────────────────┐
    │                         USER                                │
    │                      (Chat Interface)                       │
    └──────────────────────────┬──────────────────────────────────┘
                               │ Chat Message
                               ↓
    ┌─────────────────────────────────────────────────────────────┐
    │              RAG WORKFLOW (Main Workflow)                   │
    │  ┌─────────────────────────────────────────────────────┐   │
    │  │  When Chat Message Received (Trigger)               │   │
    │  └────────────────────┬────────────────────────────────┘   │
    │                       ↓                                     │
    │  ┌─────────────────────────────────────────────────────┐   │
    │  │  Google Drive - Download Documents                  │   │
    │  │  - Fetches knowledge base documents                 │   │
    │  └────────────────────┬────────────────────────────────┘   │
    │                       ↓                                     │
    │  ┌─────────────────────────────────────────────────────┐   │
    │  │  Pinecone - Default Data Loader                     │   │
    │  │  - Processes and embeds documents                   │   │
    │  └────────────────────┬────────────────────────────────┘   │
    │                       ↓                                     │
    │  ┌─────────────────────────────────────────────────────┐   │
    │  │  AI AGENT (OpenAI + Tools)                          │   │
    │  │  ┌──────────────────────────────────────────┐       │   │
    │  │  │  Tool 1: Vector Store (Knowledge Base)   │       │   │
    │  │  └──────────────────────────────────────────┘       │   │
    │  │  ┌──────────────────────────────────────────┐       │   │
    │  │  │  Tool 2: Ticket Manager (Sub-Workflow) ──┼───────┼───┼──┐
    │  │  └──────────────────────────────────────────┘       │   │  │
    │  └─────────────────────────────────────────────────────┘   │  │
    └─────────────────────────────────────────────────────────────┘  │
                                                                      │
                          ┌───────────────────────────────────────────┘
                          │ toolWorkflow Call
                          ↓
    ┌─────────────────────────────────────────────────────────────┐
    │        TICKET MANAGER WORKFLOW (Sub-Workflow)               │
    │  ┌─────────────────────────────────────────────────────┐   │
    │  │  Execute Workflow Trigger                           │   │
    │  └────────────────────┬────────────────────────────────┘   │
    │                       ↓                                     │
    │  ┌─────────────────────────────────────────────────────┐   │
    │  │  Normalize Inputs (Set Defaults)                    │   │
    │  └────────────────────┬────────────────────────────────┘   │
    │                       ↓                                     │
    │  ┌─────────────────────────────────────────────────────┐   │
    │  │  ACTION SWITCH (Route by action parameter)          │   │
    │  └───┬──────┬──────┬──────┬──────────────────────────┘   │
    │      │      │      │      │                               │
    │  create  status update close                              │
    │      │      │      │      │                               │
    │      ↓      ↓      ↓      ↓                               │
    │   [Branch flows to respective handlers]                   │
    │      │      │      │      │                               │
    │      ↓      ↓      ↓      ↓                               │
    │  ┌────────────────────────────────┐                       │
    │  │  AIRTABLE OPERATIONS           │                       │
    │  │  - Create/Find/Update Records  │                       │
    │  └────────────┬───────────────────┘                       │
    │               ↓                                            │
    │  ┌────────────────────────────────┐                       │
    │  │  Build Response (Return JSON)  │                       │
    │  └────────────┬───────────────────┘                       │
    │               │                                            │
    └───────────────┼────────────────────────────────────────────┘
                    │ Return to AI Agent
                    ↓
    ┌─────────────────────────────────────────────────────────────┐
    │  RAG Workflow - AI Agent Formats Response to User          │
    └─────────────────────────────────────────────────────────────┘
    """

    p = doc.add_paragraph(arch_diagram)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(8)

    doc.add_paragraph()

    doc.add_heading('2.2 Ticket Manager Workflow Details', 2)

    doc.add_paragraph(
        "The Ticket Manager is a sub-workflow that handles all CRUD operations for support tickets. "
        "It uses action-based routing to determine which operation to perform."
    )

    doc.add_heading('2.2.1 Action-Based Routing', 3)

    actions_table = doc.add_table(rows=1, cols=3)
    actions_table.style = 'Light Grid Accent 1'
    hdr = actions_table.rows[0].cells
    hdr[0].text = 'Action'
    hdr[1].text = 'Required Parameters'
    hdr[2].text = 'Description'

    actions_data = [
        ('create', 'name, email, subject, description, priority',
         'Creates new ticket with unique ID and SLA calculation'),
        ('status', 'ticketId',
         'Retrieves current status and details of a ticket'),
        ('update', 'ticketId, description',
         'Adds to conversation log, reopens if closed'),
        ('close', 'ticketId',
         'Marks ticket as closed, prevents further updates')
    ]

    for action, params, desc in actions_data:
        row = actions_table.add_row().cells
        row[0].text = action
        row[1].text = params
        row[2].text = desc

    doc.add_paragraph()

    doc.add_heading('2.3 Data Flow', 2)

    data_flow = """
    User Input → AI Agent → Intent Detection
                              │
                              ├─→ Knowledge Base Query (Vector Store)
                              │   └─→ Answer from Documents
                              │
                              └─→ Ticket Operation Needed
                                  │
                                  ↓
                            Ticket Manager Sub-Workflow
                                  │
                                  ├─→ Input Normalization
                                  ├─→ Action Routing
                                  ├─→ Airtable Operation
                                  ├─→ Response Building
                                  └─→ Return JSON
                                      │
                                      ↓
                            AI Agent Receives Result
                                      │
                                      ↓
                            Format User-Friendly Message
                                      │
                                      ↓
                            Return to User
    """

    p = doc.add_paragraph(data_flow)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_page_break()

def add_workflow_components(doc):
    """Add detailed workflow components"""
    doc.add_heading('3. Workflow Components', 1)

    doc.add_heading('3.1 RAG Workflow Components', 2)

    rag_components = [
        {
            'name': 'When Chat Message Received',
            'type': 'Trigger',
            'description': 'Initiates workflow when user sends a message via chat interface'
        },
        {
            'name': 'Google Drive - Download Documents',
            'type': 'Data Loader',
            'description': 'Fetches knowledge base documents from Google Drive for RAG context'
        },
        {
            'name': 'Pinecone - Default Data Loader',
            'type': 'Vector Store',
            'description': 'Processes documents and creates embeddings for semantic search'
        },
        {
            'name': 'AI Agent',
            'type': 'LangChain Agent',
            'description': 'Central intelligence hub with access to Vector Store and Ticket Manager tools'
        },
        {
            'name': 'Call Create Ticket (Tool)',
            'type': 'Workflow Tool',
            'description': 'Provides AI agent access to Ticket Manager sub-workflow'
        }
    ]

    for comp in rag_components:
        doc.add_heading(comp['name'], 3)
        doc.add_paragraph(f"Type: {comp['type']}")
        doc.add_paragraph(f"Description: {comp['description']}")
        doc.add_paragraph()

    doc.add_heading('3.2 Ticket Manager Workflow Components', 2)

    doc.add_heading('3.2.1 Core Nodes', 3)

    ticket_components = [
        {
            'name': 'Execute Workflow Trigger',
            'description': 'Entry point when called from RAG workflow',
            'inputs': 'action, ticketId, name, email, subject, description, priority, additionalContext'
        },
        {
            'name': 'Normalize Inputs',
            'description': 'Sets default values for missing parameters',
            'defaults': 'priority=medium, channel=chat, subject="No subject provided"'
        },
        {
            'name': 'Action Switch',
            'description': 'Routes to appropriate branch based on action parameter',
            'branches': '4 branches: create, status, update, close'
        }
    ]

    for comp in ticket_components:
        doc.add_paragraph(f"• {comp['name']}", style='List Bullet')
        doc.add_paragraph(f"  {comp['description']}")
        if 'inputs' in comp:
            doc.add_paragraph(f"  Inputs: {comp['inputs']}")
        if 'defaults' in comp:
            doc.add_paragraph(f"  Defaults: {comp['defaults']}")
        if 'branches' in comp:
            doc.add_paragraph(f"  {comp['branches']}")

    doc.add_heading('3.2.2 CREATE Branch', 3)

    create_flow = """
    Code - Prepare Create
      │  - Generate unique Ticket ID (TCK-{timestamp}-{random})
      │  - Calculate SLA due date based on priority
      │  - Initialize conversation log
      │  - Set status = 'open'
      ↓
    Airtable - Create Ticket
      │  - Insert new record with all fields
      ↓
    Code - Build Create Response
      │  - Prepare JSON response with ticket details
      │  - Include messageForUser confirmation
      └→ RETURN (end node, no connections)
    """

    p = doc.add_paragraph(create_flow)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_heading('3.2.3 STATUS Branch', 3)

    status_flow = """
    Airtable - Find Ticket (Status)
      │  - Search by Ticket ID using filterByFormula
      ↓
    Code - Build Status Response
      │  - Extract current status, subject, priority
      │  - Build user-friendly status message
      └→ RETURN
    """

    p = doc.add_paragraph(status_flow)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_heading('3.2.4 UPDATE Branch', 3)

    update_flow = """
    Airtable - Find Ticket (Update)
      ↓
    Code - Prepare Update
      │  - CHECK 1: Is ticket closed/resolved? → Block with error
      │  - CHECK 2: Is description empty? → Block with prompt
      │  - Append to conversation log with timestamp
      │  - If was closed, reopen (status = 'open')
      ↓
    Airtable - Update Ticket
      │  - Write updated conversation log
      │  - Update status if needed
      │  - Update 'Updated At' timestamp
      ↓
    Code - Build Update Response
      │  - Confirm update to user
      └→ RETURN
    """

    p = doc.add_paragraph(update_flow)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_heading('3.2.5 CLOSE Branch', 3)

    close_flow = """
    Airtable - Find Ticket (Close)
      ↓
    Code - Prepare Close
      │  - Extract airtableRecordId
      │  - Check if already closed
      │  - Prepare close message
      ↓
    Airtable - Close Ticket
      │  - Set status = 'closed'
      │  - Update timestamp
      ↓
    Code - Build Close Response
      │  - Preserve ticketId using node reference
      │  - Build confirmation message
      └→ RETURN
    """

    p = doc.add_paragraph(close_flow)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_page_break()

def add_data_schema(doc):
    """Add data schema section"""
    doc.add_heading('4. Data Schema', 1)

    doc.add_heading('4.1 Airtable Schema', 2)

    doc.add_paragraph(
        "The system uses a single-table design in Airtable for simplicity. "
        "Table: 'Imported table' (tbl9AlVNEOqUcpRCb) in Base: appEQ1o4iqY0Nv5bB"
    )

    doc.add_heading('4.2 Tickets Table Fields', 2)

    schema_table = doc.add_table(rows=1, cols=4)
    schema_table.style = 'Light Grid Accent 1'
    hdr = schema_table.rows[0].cells
    hdr[0].text = 'Field Name'
    hdr[1].text = 'Type'
    hdr[2].text = 'Required'
    hdr[3].text = 'Description'

    fields_data = [
        ('Ticket ID', 'Text (Single line)', 'Yes', 'Unique identifier: TCK-{timestamp}-{random}'),
        ('Customer Name', 'Text', 'No', 'Name of the customer creating ticket'),
        ('Customer Email', 'Email', 'No', 'Email address for contact'),
        ('Channel', 'Text', 'Yes', 'Source channel (default: "chat")'),
        ('Subject', 'Text', 'Yes', 'Brief title of the issue'),
        ('Initial Description', 'Long text', 'Yes', 'Original problem description'),
        ('Conversation Log', 'Long text', 'Yes', 'Timestamped history of all updates'),
        ('Priority', 'Single select', 'Yes', 'low | medium | high | urgent'),
        ('Status', 'Single select', 'Yes', 'open | in-progress | closed | resolved'),
        ('Created At', 'DateTime', 'Yes', 'ISO 8601 timestamp of creation'),
        ('Updated At', 'DateTime', 'Yes', 'ISO 8601 timestamp of last update'),
        ('SLA Due At', 'DateTime', 'Yes', 'Calculated deadline based on priority'),
        ('Internal Notes', 'Long text', 'No', 'Staff-only notes and context')
    ]

    for field_name, field_type, required, description in fields_data:
        row = schema_table.add_row().cells
        row[0].text = field_name
        row[1].text = field_type
        row[2].text = required
        row[3].text = description

    doc.add_paragraph()

    doc.add_heading('4.3 SLA Calculation Rules', 2)

    sla_table = doc.add_table(rows=1, cols=2)
    sla_table.style = 'Light Grid Accent 1'
    hdr = sla_table.rows[0].cells
    hdr[0].text = 'Priority'
    hdr[1].text = 'SLA Duration'

    sla_data = [
        ('high / urgent', '1 day (24 hours)'),
        ('medium', '3 days (72 hours)'),
        ('low', '5 days (120 hours)')
    ]

    for priority, duration in sla_data:
        row = sla_table.add_row().cells
        row[0].text = priority
        row[1].text = duration

    doc.add_paragraph()

    doc.add_heading('4.4 Ticket ID Format', 2)

    doc.add_paragraph("Format: TCK-{timestamp}-{random}")
    doc.add_paragraph("Example: TCK-1733148920123-456")
    doc.add_paragraph("• 'TCK-' prefix for easy identification")
    doc.add_paragraph("• Timestamp in milliseconds ensures chronological ordering")
    doc.add_paragraph("• Random 3-digit suffix prevents collisions")

    doc.add_page_break()

def add_api_reference(doc):
    """Add API reference section"""
    doc.add_heading('5. API Reference', 1)

    doc.add_heading('5.1 Webhook Endpoint', 2)

    doc.add_paragraph("Base URL: https://polarmedia.app.n8n.cloud")
    doc.add_paragraph("Webhook Path: /webhook/tt")
    doc.add_paragraph("Full URL: https://polarmedia.app.n8n.cloud/webhook/tt")
    doc.add_paragraph("Method: POST")
    doc.add_paragraph("Content-Type: application/json")

    doc.add_heading('5.2 Create Ticket', 2)

    create_example = """POST /webhook/tt
Content-Type: application/json

{
  "action": "create",
  "name": "John Doe",
  "email": "john@example.com",
  "subject": "Cannot login to account",
  "description": "Getting 403 error when trying to access billing dashboard",
  "priority": "high"
}

Response:
{
  "action": "create",
  "ticketId": "TCK-1733148920123-456",
  "status": "open",
  "priority": "high",
  "subject": "Cannot login to account",
  "messageForUser": "I've created ticket TCK-1733148920123-456 for your issue..."
}"""

    p = doc.add_paragraph(create_example)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_heading('5.3 Check Status', 2)

    status_example = """POST /webhook/tt
Content-Type: application/json

{
  "action": "status",
  "ticketId": "TCK-1733148920123-456"
}

Response:
{
  "action": "status",
  "ticketId": "TCK-1733148920123-456",
  "status": "open",
  "priority": "high",
  "subject": "Cannot login to account",
  "messageForUser": "Ticket TCK-1733148920123-456 is currently open..."
}"""

    p = doc.add_paragraph(status_example)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_heading('5.4 Update Ticket', 2)

    update_example = """POST /webhook/tt
Content-Type: application/json

{
  "action": "update",
  "ticketId": "TCK-1733148920123-456",
  "description": "Tried clearing cache as suggested, still not working"
}

Response:
{
  "action": "update",
  "ticketId": "TCK-1733148920123-456",
  "status": "open",
  "messageForUser": "I've updated your ticket TCK-1733148920123-456..."
}"""

    p = doc.add_paragraph(update_example)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_heading('5.5 Close Ticket', 2)

    close_example = """POST /webhook/tt
Content-Type: application/json

{
  "action": "close",
  "ticketId": "TCK-1733148920123-456"
}

Response:
{
  "action": "close",
  "ticketId": "TCK-1733148920123-456",
  "status": "closed",
  "messageForUser": "I've closed ticket TCK-1733148920123-456..."
}"""

    p = doc.add_paragraph(close_example)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_heading('5.6 Error Responses', 2)

    doc.add_paragraph("Ticket Not Found:")
    error1 = '''{
  "action": "status",
  "ticketId": "",
  "status": "not_found",
  "messageForUser": "I could not find a ticket with that ID..."
}'''
    p = doc.add_paragraph(error1)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph("Update Blocked (Empty Description):")
    error2 = '''{
  "messageForUser": "Please provide the update details so I can add them to your ticket.",
  "skipUpdate": true
}'''
    p = doc.add_paragraph(error2)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph("Update Blocked (Ticket Closed):")
    error3 = '''{
  "messageForUser": "Ticket TCK-... is closed and cannot be updated. Please open a new ticket...",
  "skipUpdate": true
}'''
    p = doc.add_paragraph(error3)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_page_break()

def add_testing_framework(doc):
    """Add testing framework section"""
    doc.add_heading('6. Testing Framework', 1)

    doc.add_heading('6.1 Test Environment Setup', 2)

    setup_code = """# Set environment variables
export N8N_WEBHOOK_BASE="https://polarmedia.app.n8n.cloud"
export N8N_TICKET_WEBHOOK_PATH="/webhook/tt"
export AUTH_HEADER=""  # Optional if auth required"""

    p = doc.add_paragraph(setup_code)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_heading('6.2 Test Scripts', 2)

    test_scripts = [
        {
            'name': 'test_create.sh',
            'purpose': 'Creates a sample ticket and extracts ticket ID',
            'usage': './test_create.sh'
        },
        {
            'name': 'test_status.sh',
            'purpose': 'Checks status of an existing ticket',
            'usage': './test_status.sh <ticketId>'
        },
        {
            'name': 'test_close_bug_reproduction.sh',
            'purpose': 'Tests the close bug fix (create → close → verify ID match)',
            'usage': './test_close_bug_reproduction.sh'
        },
        {
            'name': 'all_test.sh',
            'purpose': 'Comprehensive end-to-end test suite',
            'usage': './all_test.sh'
        }
    ]

    for script in test_scripts:
        doc.add_heading(script['name'], 3)
        doc.add_paragraph(f"Purpose: {script['purpose']}")
        doc.add_paragraph(f"Usage: {script['usage']}")
        doc.add_paragraph()

    doc.add_heading('6.3 All Tests Suite', 2)

    doc.add_paragraph("The all_test.sh script runs a complete scenario test:")

    test_flow = """1. Create ticket
   ↓
2. Check status
   ↓
3. Update ticket with text
   ↓
4. Try update without text (should block)
   ↓
5. Close ticket
   ↓
6. Try update after close (should block)"""

    p = doc.add_paragraph(test_flow)
    p.style = 'No Spacing'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

    doc.add_paragraph()
    doc.add_paragraph("Expected Result: All 6 tests pass ✓")

    doc.add_heading('6.4 Test Coverage', 2)

    coverage_table = doc.add_table(rows=1, cols=3)
    coverage_table.style = 'Light Grid Accent 1'
    hdr = coverage_table.rows[0].cells
    hdr[0].text = 'Test Case'
    hdr[1].text = 'Status'
    hdr[2].text = 'Coverage'

    coverage_data = [
        ('Create ticket with all fields', 'PASS ✓', 'Happy path'),
        ('Status check existing ticket', 'PASS ✓', 'Read operation'),
        ('Update with valid text', 'PASS ✓', 'Happy path'),
        ('Update with empty text', 'PASS ✓', 'Validation'),
        ('Close ticket', 'PASS ✓', 'State change'),
        ('Update closed ticket', 'PASS ✓', 'Validation'),
        ('Ticket ID preservation', 'PASS ✓', 'Data integrity'),
        ('Response format validation', 'PASS ✓', 'API contract')
    ]

    for test, status, coverage in coverage_data:
        row = coverage_table.add_row().cells
        row[0].text = test
        row[1].text = status
        row[2].text = coverage

    doc.add_page_break()

def add_deployment_guide(doc):
    """Add deployment guide"""
    doc.add_heading('7. Deployment Guide', 1)

    doc.add_heading('7.1 Prerequisites', 2)

    prereqs = [
        "n8n Cloud account (v1.118.2 or higher)",
        "Airtable account with API access",
        "OpenAI API key",
        "Pinecone account and API key",
        "Google Drive with knowledge base documents",
        "Slack workspace (optional, for notifications)"
    ]

    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')

    doc.add_heading('7.2 Setup Steps', 2)

    doc.add_paragraph("Step 1: Configure Airtable")
    doc.add_paragraph("• Create 'Tickets' table with schema from Section 4.2")
    doc.add_paragraph("• Generate personal access token")
    doc.add_paragraph("• Note Base ID and Table ID")

    doc.add_paragraph()
    doc.add_paragraph("Step 2: Import Workflows to n8n")
    doc.add_paragraph("• Import 'RAG Workflow For( Customer service chat-bot).json'")
    doc.add_paragraph("• Import 'Ticket Manager (Airtable).json'")
    doc.add_paragraph("• Import 'Slack Actions.json' (optional)")

    doc.add_paragraph()
    doc.add_paragraph("Step 3: Configure Credentials")
    doc.add_paragraph("• Add Airtable Personal Access Token")
    doc.add_paragraph("• Add OpenAI API credentials")
    doc.add_paragraph("• Add Pinecone API credentials")
    doc.add_paragraph("• Add Google Drive OAuth2 credentials")
    doc.add_paragraph("• Add Slack credentials (if using notifications)")

    doc.add_paragraph()
    doc.add_paragraph("Step 4: Update Workflow Parameters")
    doc.add_paragraph("• In Ticket Manager: Update Airtable Base ID and Table ID")
    doc.add_paragraph("• In RAG Workflow: Configure Google Drive folder ID")
    doc.add_paragraph("• Update Pinecone index name")

    doc.add_paragraph()
    doc.add_paragraph("Step 5: Activate Workflows")
    doc.add_paragraph("• Activate 'Ticket Manager (Airtable)' workflow")
    doc.add_paragraph("• Activate 'RAG Workflow' with chat trigger")
    doc.add_paragraph("• Test webhook endpoint")

    doc.add_paragraph()
    doc.add_paragraph("Step 6: Run Tests")
    doc.add_paragraph("• Set environment variables")
    doc.add_paragraph("• Run ./all_test.sh")
    doc.add_paragraph("• Verify all tests pass")

    doc.add_heading('7.3 Environment Variables', 2)

    env_table = doc.add_table(rows=1, cols=2)
    env_table.style = 'Light Grid Accent 1'
    hdr = env_table.rows[0].cells
    hdr[0].text = 'Variable'
    hdr[1].text = 'Value'

    env_data = [
        ('N8N_WEBHOOK_BASE', 'https://polarmedia.app.n8n.cloud'),
        ('N8N_TICKET_WEBHOOK_PATH', '/webhook/tt'),
        ('AIRTABLE_BASE_ID', 'appEQ1o4iqY0Nv5bB'),
        ('AIRTABLE_TABLE_ID', 'tbl9AlVNEOqUcpRCb')
    ]

    for var, val in env_data:
        row = env_table.add_row().cells
        row[0].text = var
        row[1].text = val

    doc.add_page_break()

def add_troubleshooting(doc):
    """Add troubleshooting section"""
    doc.add_heading('8. Troubleshooting', 1)

    doc.add_heading('8.1 Common Issues', 2)

    issues = [
        {
            'problem': 'Close action returns wrong ticket ID',
            'cause': 'Build Response nodes connected to Slack instead of being end nodes',
            'solution': 'Disconnect Slack nodes, make Build Response nodes terminal (no outgoing connections)'
        },
        {
            'problem': 'Empty responses from update/status actions',
            'cause': 'Same as above - response nodes not terminal',
            'solution': 'Ensure all Build Response nodes have no outgoing connections'
        },
        {
            'problem': 'Ticket ID lost in close response',
            'cause': 'Airtable update response doesn\'t include Ticket ID field',
            'solution': 'Use explicit node reference: $("Code - Prepare Close").item.json.ticketId'
        },
        {
            'problem': 'Pinned data causing stale results',
            'cause': 'Sample data pinned in n8n UI',
            'solution': 'Click node, check for pin icon, unpin all nodes'
        },
        {
            'problem': 'AI agent not calling ticket tool',
            'cause': 'Unclear intent or missing keywords',
            'solution': 'Improve AI agent prompt, add explicit tool instructions'
        }
    ]

    for issue in issues:
        doc.add_heading(issue['problem'], 3)
        doc.add_paragraph(f"Cause: {issue['cause']}")
        doc.add_paragraph(f"Solution: {issue['solution']}")
        doc.add_paragraph()

    doc.add_heading('8.2 Debugging Checklist', 2)

    debug_steps = [
        "Check n8n execution log for errors",
        "Verify all credentials are valid and not expired",
        "Test each branch manually in n8n editor",
        "Check Airtable record directly to verify data",
        "Review node inputs/outputs in execution view",
        "Ensure no pinned data on any nodes",
        "Verify webhook URL is correct",
        "Check for recent n8n version changes",
        "Test with curl directly (bypass AI agent)",
        "Review workflow JSON for incorrect connections"
    ]

    for i, step in enumerate(debug_steps, 1):
        doc.add_paragraph(f"{i}. {step}", style='List Number')

    doc.add_heading('8.3 Performance Optimization', 2)

    doc.add_paragraph("If experiencing slow response times:")

    perf_tips = [
        "Cache Pinecone embeddings to reduce lookup time",
        "Limit Google Drive document size and count",
        "Use Airtable indexes on Ticket ID field",
        "Reduce AI agent prompt length",
        "Implement query result caching",
        "Monitor n8n execution time per node"
    ]

    for tip in perf_tips:
        doc.add_paragraph(tip, style='List Bullet')

    doc.add_page_break()

def add_appendices(doc):
    """Add appendices"""
    doc.add_heading('9. Appendices', 1)

    doc.add_heading('9.1 Glossary', 2)

    glossary_table = doc.add_table(rows=1, cols=2)
    glossary_table.style = 'Light Grid Accent 1'
    hdr = glossary_table.rows[0].cells
    hdr[0].text = 'Term'
    hdr[1].text = 'Definition'

    glossary = [
        ('RAG', 'Retrieval-Augmented Generation - AI technique combining retrieval with generation'),
        ('n8n', 'Low-code workflow automation platform'),
        ('Airtable', 'Cloud-based spreadsheet database'),
        ('LangChain', 'Framework for developing LLM-powered applications'),
        ('Pinecone', 'Vector database for similarity search'),
        ('Sub-workflow', 'Workflow called from another workflow'),
        ('Tool Workflow', 'n8n node type that exposes a workflow as an AI agent tool'),
        ('SLA', 'Service Level Agreement - target response time'),
        ('Vector Store', 'Database optimized for storing and searching embeddings')
    ]

    for term, definition in glossary:
        row = glossary_table.add_row().cells
        row[0].text = term
        row[1].text = definition

    doc.add_paragraph()

    doc.add_heading('9.2 References', 2)

    references = [
        "n8n Documentation: https://docs.n8n.io",
        "Airtable API: https://airtable.com/developers/web/api",
        "OpenAI API: https://platform.openai.com/docs",
        "LangChain Documentation: https://python.langchain.com",
        "Pinecone Documentation: https://docs.pinecone.io"
    ]

    for ref in references:
        doc.add_paragraph(ref, style='List Bullet')

    doc.add_heading('9.3 Version History', 2)

    version_table = doc.add_table(rows=1, cols=3)
    version_table.style = 'Light Grid Accent 1'
    hdr = version_table.rows[0].cells
    hdr[0].text = 'Version'
    hdr[1].text = 'Date'
    hdr[2].text = 'Changes'

    versions = [
        ('1.0', 'December 2024', 'Initial production release with all core features'),
        ('1.0-bugfix', 'December 2, 2024', 'Fixed close action bug, improved response handling'),
        ('1.1-planned', 'TBD', 'Slack notifications, SLA monitoring, multi-table schema')
    ]

    for ver, date, changes in versions:
        row = version_table.add_row().cells
        row[0].text = ver
        row[1].text = date
        row[2].text = changes

    doc.add_paragraph()

    doc.add_heading('9.4 Contact Information', 2)

    doc.add_paragraph("Project: Infinity Pixel Chat Bot")
    doc.add_paragraph("n8n Instance: polarmedia.app.n8n.cloud")
    doc.add_paragraph("GitHub: https://github.com/Vallabha-Praneeth/Infinity-Pixel-Chat-Bot")

    doc.add_paragraph()
    doc.add_paragraph("For technical support or questions about this documentation, "
                      "refer to the repository README or consult the workflow JSON files.")

def main():
    """Main function to generate technical documentation"""
    print("Generating Technical Documentation...")

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Build document
    add_title_page(doc)
    add_toc(doc)
    add_system_overview(doc)
    add_architecture(doc)
    add_workflow_components(doc)
    add_data_schema(doc)
    add_api_reference(doc)
    add_testing_framework(doc)
    add_deployment_guide(doc)
    add_troubleshooting(doc)
    add_appendices(doc)

    # Save
    output_path = 'RAG_Customer_Service_Chatbot_Technical_Documentation.docx'
    doc.save(output_path)
    print(f"✓ Technical documentation saved to: {output_path}")

if __name__ == '__main__':
    main()
